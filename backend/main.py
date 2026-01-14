# Import torch first to initialize CUDA DLLs
import torch
import os

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

# Standard Imports
import shutil
import uvicorn
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
import sys
import types

# Shim patch
try:
    import langchain_core.documents
    sys.modules["langchain.docstore.document"] = types.ModuleType("langchain.docstore.document")
    sys.modules["langchain.docstore.document"].Document = langchain_core.documents.Document
except: pass

# Custom Modules
from ingestion import ingestion_engine
from vector_store import get_vector_store
from agent_engine import agent_app

app = FastAPI(title="Omni-Scribe API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize Database
try:
    vector_db = get_vector_store()
except Exception as e:
    print(f"⚠️ Vector DB Init Warning: {e}")

@app.get("/")
def health_check():
    return {"status": "Omni-Scribe Brain is Active 🧠"}

@app.post("/ingest/audio")
async def ingest_audio(file: UploadFile = File(...)):
    temp_path = f"temp_{file.filename}"
    try:
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        print(f"🎤 Transcribing {file.filename}...")
        text = ingestion_engine.transcribe_audio(temp_path)
        
        vector_db.add_texts(texts=[text], metadatas=[{"source": "audio", "filename": file.filename}])
        return {"status": "success", "text_snippet": text[:100] + "..."}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@app.post("/ingest/image")
async def ingest_image(file: UploadFile = File(...)):
    temp_path = f"temp_{file.filename}"
    try:
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        print(f"👁️ Analyzing Image {file.filename}...")
        text = ingestion_engine.extract_text_from_image(temp_path)
        
        vector_db.add_texts(texts=[text], metadatas=[{"source": "image", "filename": file.filename}])
        return {"status": "success", "extracted_text": text[:100] + "..."}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@app.post("/chat")
async def chat(query: str = Form(...)):
    print(f"💬 Processing Query: {query}")
    inputs = {"query": query, "context": [], "response": "", "is_sufficient": False}
    
    # Trigger LangGraph workflow
    result = agent_app.invoke(inputs)

    print(f"🤖 AI Response: {result['response']}")
    print("-" * 50)
    
    return {
        "answer": result["response"],
        "context_used": result["context"]
    }

@app.post("/feedback")
async def learn_from_feedback(
    original_query: str = Form(...), 
    correct_answer: str = Form(...)
):
    """
    Implements the Self-Learning Memory System.
    Ingests human corrections as 'High Priority' context.
    """
    print(f"🧠 Self-Learning triggered for: {original_query}")
    
    # Create a synthetic document with the correction
    learning_text = f"[HUMAN CORRECTION] Question: {original_query}\nAnswer: {correct_answer}"
    
    # Save to Vector DB immediately
    vector_db.add_texts(
        texts=[learning_text],
        metadatas=[{"source": "human_feedback", "type": "correction"}]
    )
    
    return {"status": "learned", "message": "Memory updated. I won't make that mistake again."}

@app.post("/ingest/text")
async def ingest_text(file: UploadFile = File(...)):
    """
    Ingest document files (.txt, .md, .pdf, .docx) into the knowledge base.
    Splits large files into chunks for better retrieval.
    """
    # Validate file type
    allowed_extensions = ['.txt', '.md', '.pdf', '.docx']
    file_ext = os.path.splitext(file.filename)[1].lower()
    if file_ext not in allowed_extensions:
        raise HTTPException(status_code=400, detail=f"Only {allowed_extensions} files are supported")
    
    try:
        # Read file content based on type
        content = await file.read()
        
        if file_ext in ['.txt', '.md']:
            text = content.decode('utf-8')
        elif file_ext == '.pdf':
            from pypdf import PdfReader
            import io
            pdf_reader = PdfReader(io.BytesIO(content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        elif file_ext == '.docx':
            from docx import Document
            import io
            doc = Document(io.BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        
        print(f"📄 Processing document: {file.filename} ({len(text)} chars)")
        
        # Split into chunks if large (>1000 chars per chunk)
        chunk_size = 1000
        chunks = []
        if len(text) > chunk_size:
            # Split by paragraphs first, then by size
            paragraphs = text.split('\n\n')
            current_chunk = ""
            for para in paragraphs:
                if len(current_chunk) + len(para) < chunk_size:
                    current_chunk += para + "\n\n"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para + "\n\n"
            if current_chunk:
                chunks.append(current_chunk.strip())
        else:
            chunks = [text]
        
        # Add to vector DB
        for i, chunk in enumerate(chunks):
            tagged_text = f"[DOCUMENT - {file_ext.upper()}]: {chunk}"
            vector_db.add_texts(
                texts=[tagged_text], 
                metadatas=[{"source": "document", "type": file_ext, "filename": file.filename, "chunk": i}]
            )
        
        return {
            "status": "success", 
            "filename": file.filename,
            "file_type": file_ext,
            "chunks_created": len(chunks),
            "text_snippet": text[:100] + "..." if len(text) > 100 else text
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/ingest/scan")
async def scan_knowledge_folder():
    """
    Scan the knowledge folder and ingest all document files.
    Supports: .txt, .md, .pdf, .docx
    Directory: D:\\AIML-Projects\\OmniScribe\\knowledge
    """
    import config
    import glob
    
    knowledge_dir = config.KNOWLEDGE_DIR
    
    if not os.path.exists(knowledge_dir):
        os.makedirs(knowledge_dir)
        return {"status": "created", "message": f"Created empty knowledge folder at {knowledge_dir}", "files_processed": 0}
    
    # Find all document files
    patterns = ['*.txt', '*.md', '*.pdf', '*.docx']
    files_found = []
    for pattern in patterns:
        files_found.extend(glob.glob(os.path.join(knowledge_dir, pattern)))
    
    if not files_found:
        return {"status": "empty", "message": "No document files found in knowledge folder", "files_processed": 0}
    
    processed = []
    errors = []
    
    for filepath in files_found:
        try:
            filename = os.path.basename(filepath)
            file_ext = os.path.splitext(filename)[1].lower()
            
            # Extract text based on file type
            if file_ext in ['.txt', '.md']:
                with open(filepath, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_ext == '.pdf':
                from pypdf import PdfReader
                pdf_reader = PdfReader(filepath)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            elif file_ext == '.docx':
                from docx import Document
                doc = Document(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                continue
            
            print(f"📂 Scanning: {filename} ({len(text)} chars)")
            
            # Add to vector DB with file tag
            tagged_text = f"[DOCUMENT - {filename}]: {text}"
            vector_db.add_texts(
                texts=[tagged_text], 
                metadatas=[{"source": "knowledge_folder", "type": file_ext, "filename": filename, "path": filepath}]
            )
            processed.append(filename)
            
        except Exception as e:
            errors.append({"file": filename, "error": str(e)})
    
    return {
        "status": "success",
        "files_processed": len(processed),
        "files": processed,
        "errors": errors if errors else None
    }

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
